import os
import json
import uuid
import io
from datetime import datetime
from urllib.parse import quote
from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import StreamingResponse
from schemas import TopicGenerateRequest, TopicGenerateResponse
from services.prompt_engine import prompt_engine
from services.srs import SRSEngine
from db.database import get_db
from db.auth_db import get_user_profile

router = APIRouter(prefix="/topic", tags=["话题生成"])


@router.post(
    "/generate",
    response_model=TopicGenerateResponse,
    summary="生成话题文章",
    description="根据指定主题、领域与难度生成学习文章，并返回新词与术语。",
)
async def generate_topic(req: TopicGenerateRequest, request: Request, db=Depends(get_db)):
    # V2.0: 读取用户画像，注入专业方向到提示词
    user_profile = {}
    try:
        user_profile = get_user_profile(request.state.user["id"])
    except Exception:
        pass

    # SRS 引擎：获取需要复习的目标词汇
    srs = SRSEngine(db)
    db_words = req.db_words
    if not db_words and req.db_word_count > 0:
        db_words = srs.select_words_for_topic(limit=req.db_word_count)

    # 从数据库读取文章风格的 prompt_modifier
    style_modifier = ""
    style_row = db.execute(
        "SELECT prompt_modifier FROM article_styles WHERE id=?",
        (req.article_style_id,),
    ).fetchone()
    if style_row:
        style_modifier = style_row["prompt_modifier"] or ""

    # V2.0: 获取用户已有生词本词汇，用于排除
    existing_vocab = []
    try:
        rows = db.execute("SELECT word FROM vocab").fetchall()
        existing_vocab = [r["word"].lower() for r in rows]
    except Exception:
        pass

    # V2.2: 获取用户最近文章摘要，用于防重复（≤30%）
    previous_articles = []
    try:
        recent_rows = db.execute(
            "SELECT topic, substr(result_text, 1, 300) AS preview FROM article_history ORDER BY created_at DESC LIMIT 10"
        ).fetchall()
        for rr in recent_rows:
            previous_articles.append(f"Topic: {rr['topic']}\nPreview: {rr['preview']}")
    except Exception:
        pass

    # 获取用户备考目标，用于限定新词范围
    exam_type = user_profile.get("exam_type", req.target_range_id or "cet6")

    # V2.1: 三层架构调用 — 用户画像与功能参数分离
    try:
        response = await prompt_engine.generate_with_context(
            template_name="topic_generate.j2",
            response_model=TopicGenerateResponse,
            max_tokens=int(req.article_length) * 6 + 2000,
            feature_id="topic_generate",
            # 第 2 层: 用户画像（自动注入 global_context.j2）
            user_profile={
                "user_specialty": user_profile.get("specialty", ""),
                "user_subject_field": user_profile.get("subject_field", ""),
                "user_interest_tags": user_profile.get("interest_tags", []),
                "user_exam_type": exam_type,
                "user_eng_level": user_profile.get("eng_level", ""),
            },
            feature_params={
                "topics": req.topics,
                "domains": req.domains,
                "level": req.level,
                "article_length": req.article_length,
                "style_modifier": style_modifier,
                "db_words": db_words,
                "new_word_count": req.new_word_count,
                "exam_type": exam_type,
                "existing_vocab": existing_vocab[:100],
                "previous_articles": previous_articles,
                "events": req.events or "",
            },
        )
    except Exception as e:
        import logging
        logging.getLogger("xiaotiao").error("Topic generate error: %s", e)
        raise HTTPException(
            status_code=500, detail="AI 生成失败，请稍后重试。"
        )

    # SRS 更新：记录词汇曝光
    srs.process_article_exposure("article_runtime_id", response.db_words_used)

    # ── 保存到文章历史 ──
    article_id = None
    try:
        article_id = str(uuid.uuid4())
        db.execute(
            """INSERT INTO article_history
               (id, topic, domains, level, style, article_length,
                result_text, translation_text, terms_json, new_words_json,
                notes_json, confidence_hint)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                article_id,
                ", ".join(req.topics),
                json.dumps(req.domains, ensure_ascii=False),
                req.level,
                req.article_style_id,
                req.article_length,
                response.result_text,
                response.translation_text or "",
                json.dumps([t.model_dump() for t in response.terms], ensure_ascii=False),
                json.dumps([w.model_dump() for w in response.new_words], ensure_ascii=False),
                json.dumps(response.notes, ensure_ascii=False),
                response.confidence_hint,
            ),
        )
        db.commit()
    except Exception as e:
        import logging
        logging.getLogger("xiaotiao").warning("Failed to save article history: %s", e)

    # Return response with article_id for annotation binding
    result = response.model_dump() if hasattr(response, 'model_dump') else dict(response)
    if article_id:
        result["article_id"] = article_id
    return result


# ── 文章历史 API ──────────────────────────────────

@router.get("/history", summary="获取历史文章列表")
def list_article_history(page: int = 1, size: int = 20, db=Depends(get_db)):
    offset = (max(1, page) - 1) * size
    rows = db.execute(
        """SELECT id, topic, domains, level, style, article_length,
                  substr(result_text, 1, 200) AS preview,
                  confidence_hint, created_at
           FROM article_history
           ORDER BY created_at DESC
           LIMIT ? OFFSET ?""",
        (size, offset),
    ).fetchall()

    total = db.execute("SELECT COUNT(*) FROM article_history").fetchone()[0]
    items = []
    for r in rows:
        items.append({
            "id": r["id"],
            "topic": r["topic"],
            "domains": json.loads(r["domains"]) if r["domains"] else [],
            "level": r["level"],
            "style": r["style"],
            "article_length": r["article_length"],
            "preview": r["preview"],
            "confidence_hint": r["confidence_hint"],
            "created_at": r["created_at"],
        })
    return {"items": items, "total": total, "page": page, "size": size}


@router.get("/history/{article_id}", summary="获取单篇历史文章详情")
def get_article_detail(article_id: str, db=Depends(get_db)):
    row = db.execute(
        "SELECT * FROM article_history WHERE id=?", (article_id,)
    ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="文章不存在")
    return {
        "id": row["id"],
        "topic": row["topic"],
        "domains": json.loads(row["domains"]) if row["domains"] else [],
        "level": row["level"],
        "style": row["style"],
        "article_length": row["article_length"],
        "result_text": row["result_text"],
        "translation_text": row["translation_text"],
        "terms": json.loads(row["terms_json"]) if row["terms_json"] else [],
        "new_words": json.loads(row["new_words_json"]) if row["new_words_json"] else [],
        "notes": json.loads(row["notes_json"]) if row["notes_json"] else [],
        "confidence_hint": row["confidence_hint"],
        "created_at": row["created_at"],
    }


@router.delete("/history/{article_id}", summary="删除单篇历史文章")
def delete_article(article_id: str, db=Depends(get_db)):
    cur = db.execute("DELETE FROM article_history WHERE id=?", (article_id,))
    db.commit()
    if cur.rowcount == 0:
        raise HTTPException(status_code=404, detail="文章不存在")
    return {"ok": True}


@router.get("/history/{article_id}/export", summary="导出文章为 Word 文档")
def export_article_word(article_id: str, db=Depends(get_db)):
    row = db.execute(
        "SELECT * FROM article_history WHERE id=?", (article_id,)
    ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="文章不存在")

    import re
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # ── Style setup ──
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # ── Title ──
    title_para = doc.add_heading(level=0)
    title_run = title_para.add_run(f"主题探索：{row['topic']}")
    title_run.font.size = Pt(22)

    # ── Meta info ──
    LEVEL_MAP = {"beginner": "初级", "intermediate": "中级", "advanced": "高级"}
    STYLE_MAP = {
        "economist": "The Economist",
        "guardian": "The Guardian",
        "ft": "Financial Times",
        "academic": "学术期刊",
        "plain_english": "简明日常",
    }
    domains = json.loads(row["domains"]) if row["domains"] else []
    meta_text = (
        f"方向：{', '.join(domains)}  |  "
        f"等级：{LEVEL_MAP.get(row['level'], row['level'])}  |  "
        f"风格：{STYLE_MAP.get(row['style'], row['style'])}  |  "
        f"长度：{row['article_length']} 词"
    )
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(meta_text)
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"生成时间：{row['created_at']}")
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    # ── Helper: strip HTML tags ──
    def strip_html(text):
        if not text:
            return ""
        return re.sub(r"<[^>]+>", "", text).strip()

    # ── Section 1: English Article ──
    doc.add_heading("📄 英文学习文章", level=1)
    article_text = strip_html(row["result_text"] or "")
    for paragraph in article_text.split("\n"):
        p = paragraph.strip()
        if p:
            doc.add_paragraph(p)

    # ── Section 2: Chinese Translation ──
    if row["translation_text"]:
        doc.add_heading("🀄 逐段中文翻译", level=1)
        trans_text = strip_html(row["translation_text"])
        for paragraph in trans_text.split("\n"):
            p = paragraph.strip()
            if p:
                para = doc.add_paragraph(p)
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── Section 3: Key Terms ──
    terms = json.loads(row["terms_json"]) if row["terms_json"] else []
    if terms:
        doc.add_heading(f"📚 关键术语 ({len(terms)})", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        hdr = table.rows[0].cells
        hdr[0].text = "术语"
        hdr[1].text = "中文释义"
        hdr[2].text = "例句"
        for t in terms:
            row_cells = table.add_row().cells
            row_cells[0].text = t.get("term", "")
            row_cells[1].text = t.get("zh", "")
            row_cells[2].text = t.get("example", "")

    # ── Section 4: New Words ──
    new_words = json.loads(row["new_words_json"]) if row["new_words_json"] else []
    if new_words:
        doc.add_heading(f"🆕 新词 ({len(new_words)})", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        hdr = table.rows[0].cells
        hdr[0].text = "单词"
        hdr[1].text = "中文释义"
        hdr[2].text = "文中例句"
        for w in new_words:
            row_cells = table.add_row().cells
            row_cells[0].text = w.get("word", "")
            row_cells[1].text = w.get("definition_zh", "")
            row_cells[2].text = w.get("in_sentence", "")

    # ── Section 5: Notes ──
    notes = json.loads(row["notes_json"]) if row["notes_json"] else []
    if notes:
        doc.add_heading("💡 核心概念说明", level=1)
        for note in notes:
            doc.add_paragraph(note, style="List Bullet")

    # ── Section 6: Footer ──
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run("由 小跳 · 涉外法治英语学习平台 生成")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Save to buffer and return ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_topic = re.sub(r'[\\/:*?"<>|\s]+', "_", row["topic"])[:30]
    filename = f"topic_{safe_topic}_{row['created_at'][:10]}.docx"
    encoded_filename = quote(filename)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
    )
