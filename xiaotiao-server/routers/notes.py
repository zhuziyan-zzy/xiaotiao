"""笔记系统 — 通用 CRUD + Word 导出"""
import uuid
import io
from datetime import datetime
from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.responses import StreamingResponse
from db.database import get_db

router = APIRouter(prefix="/notes", tags=["笔记"])


@router.get(
    "",
    summary="查询笔记",
    description="按模块和关联 ID 查询笔记列表。",
)
def list_notes(
    module: str = Query(..., description="模块: article|translation|vocab|paper"),
    ref_id: str = Query(..., description="关联对象 ID"),
    db=Depends(get_db),
):
    rows = db.execute(
        "SELECT * FROM notes WHERE module = ? AND ref_id = ? ORDER BY created_at DESC",
        (module, ref_id),
    ).fetchall()
    return {
        "items": [
            {
                "id": r['id'],
                "module": r['module'],
                "ref_id": r['ref_id'],
                "content": r['content'],
                "created_at": r['created_at'],
                "updated_at": r['updated_at'],
            }
            for r in rows
        ]
    }


@router.get(
    "/recent",
    summary="最近笔记",
    description="获取最近的笔记列表。",
)
def recent_notes(
    module: str = Query(None),
    limit: int = Query(20, ge=1, le=100),
    db=Depends(get_db),
):
    if module:
        rows = db.execute(
            "SELECT * FROM notes WHERE module = ? ORDER BY updated_at DESC LIMIT ?",
            (module, limit),
        ).fetchall()
    else:
        rows = db.execute(
            "SELECT * FROM notes ORDER BY updated_at DESC LIMIT ?",
            (limit,),
        ).fetchall()
    return {
        "items": [
            {
                "id": r['id'],
                "module": r['module'],
                "ref_id": r['ref_id'],
                "content": r['content'],
                "created_at": r['created_at'],
                "updated_at": r['updated_at'],
            }
            for r in rows
        ]
    }


@router.post(
    "",
    summary="创建笔记",
    description="为指定模块和对象添加一条笔记。",
)
def create_note(body: dict, db=Depends(get_db)):
    module = body.get("module")
    ref_id = body.get("ref_id")
    content = body.get("content", "").strip()
    if not module or not ref_id or not content:
        raise HTTPException(status_code=400, detail="module, ref_id, content 均为必填")
    note_id = str(uuid.uuid4())
    now = datetime.now().isoformat()
    db.execute(
        "INSERT INTO notes (id, module, ref_id, content, created_at, updated_at) VALUES (?, ?, ?, ?, ?, ?)",
        (note_id, module, ref_id, content, now, now),
    )
    db.commit()
    return {"id": note_id, "module": module, "ref_id": ref_id, "content": content, "created_at": now}


@router.put(
    "/{note_id}",
    summary="编辑笔记",
)
def update_note(note_id: str, body: dict, db=Depends(get_db)):
    content = body.get("content", "").strip()
    if not content:
        raise HTTPException(status_code=400, detail="content 不能为空")
    now = datetime.now().isoformat()
    cursor = db.execute(
        "UPDATE notes SET content = ?, updated_at = ? WHERE id = ?",
        (content, now, note_id),
    )
    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="笔记不存在")
    db.commit()
    return {"status": "ok", "id": note_id}


@router.delete(
    "/{note_id}",
    summary="删除笔记",
)
def delete_note(note_id: str, db=Depends(get_db)):
    cursor = db.execute("DELETE FROM notes WHERE id = ?", (note_id,))
    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="笔记不存在")
    db.commit()
    return {"status": "ok"}


@router.get(
    "/export",
    summary="导出笔记为 Word",
    description="按模块导出笔记为 Word 文档。",
)
def export_notes(
    module: str = Query(None),
    ref_id: str = Query(None),
    db=Depends(get_db),
):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        raise HTTPException(status_code=500, detail="服务器未安装 python-docx")

    params = []
    sql = "SELECT * FROM notes"
    conditions = []
    if module:
        conditions.append("module = ?")
        params.append(module)
    if ref_id:
        conditions.append("ref_id = ?")
        params.append(ref_id)
    if conditions:
        sql += " WHERE " + " AND ".join(conditions)
    sql += " ORDER BY created_at DESC"

    rows = db.execute(sql, params).fetchall()
    if not rows:
        raise HTTPException(status_code=404, detail="没有可导出的笔记")

    doc = Document()
    doc.add_heading("我的笔记", level=1)

    MODULE_LABELS = {"article": "文章", "translation": "翻译", "vocab": "词汇", "paper": "论文"}
    for r in rows:
        label = MODULE_LABELS.get(r['module'], r['module'])
        doc.add_heading(f"[{label}] {r['ref_id'][:8]}...", level=2)
        p = doc.add_paragraph(r['content'])
        p.style.font.size = Pt(11)
        doc.add_paragraph(f"创建: {r['created_at']}  更新: {r['updated_at']}").italic = True
        doc.add_paragraph("")  # spacer

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"notes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )
