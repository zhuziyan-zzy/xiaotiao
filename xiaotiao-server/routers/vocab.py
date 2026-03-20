import uuid
import os
import base64
import csv
import asyncio
from io import BytesIO, StringIO
from datetime import datetime, date, timedelta
from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from typing import List, Optional, Dict, Any
from db.database import get_db
from db.auth_db import get_user_profile
from schemas_vocab import VocabItemCreate, VocabItemUpdate, VocabItemResponse, VocabStatsResponse, VocabListResponse
from services.llm import call_claude_json, call_claude_vision_json
from fastapi import Request as FastAPIRequest

router = APIRouter(prefix="/vocab", tags=["生词本"])


def _load_prompt(filename: str) -> str:
    path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "prompts", filename)
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    return ""


VOCAB_IMPORT_SYSTEM_PROMPT = """You are a vocabulary extraction assistant. Given the content from a file, extract English vocabulary words suitable for a language learner's word bank.

Return a JSON object with this structure:
{
  "words": [
    {
      "word": "jurisdiction",
      "definition_zh": "管辖权",
      "part_of_speech": "n.",
      "example_sentence": "The court has jurisdiction over this matter."
    }
  ]
}

Rules:
- Extract meaningful English words/phrases (not common words like "the", "is", "a")
- Provide accurate Chinese definitions
- Part of speech should be one of: n., v., adj., adv., prep., phrase
- If the content already has word-definition pairs, preserve them
- If the content is a word list without definitions, generate appropriate definitions
- For CSV/tabular data, detect which column contains words and which contains definitions
- Maximum 100 words per extraction
- Return ONLY the JSON object, no markdown wrappers
"""


# SQLite3 compatibility shim — passes through raw SQL strings.
# Named-parameter queries (`:param`) work directly with SQLite3 execute().
def text(query: str) -> str:
    return query

@router.get(
    "",
    response_model=VocabListResponse,
    summary="获取生词列表",
    description="分页获取生词本列表，支持搜索与领域筛选。",
)
def get_vocab_list(
    db = Depends(get_db),
    page: int = Query(1, ge=1),
    limit: int = Query(50, ge=1, le=200),
    search: str = None,
    domain: str = None,
    filter: str = None,
    date: str = None
):
    offset = (page - 1) * limit
    
    # Build query
    base_query = """
        SELECT v.*, s.traversal_count, s.ease_factor, s.interval_days, 
               s.next_review_date, s.is_mastered
        FROM vocabulary_items v
        LEFT JOIN vocabulary_srs_states s ON v.id = s.vocab_id
        WHERE 1=1
    """
    count_query = "SELECT count(1) FROM vocabulary_items v WHERE 1=1"
    
    params = {}
    if search:
        base_query += " AND v.word LIKE :search"
        count_query += " AND v.word LIKE :search"
        params['search'] = f"%{search}%"
    if domain:
        base_query += " AND v.domain = :domain"
        count_query += " AND v.domain = :domain"
        params['domain'] = domain

    # Time-based and status filters
    if filter == 'today':
        target_date = date or datetime.now().strftime('%Y-%m-%d')
        base_query += " AND date(v.created_at) = :target_date"
        count_query += " AND date(v.created_at) = :target_date"
        params['target_date'] = target_date
    elif filter == 'month':
        target_month = date or datetime.now().strftime('%Y-%m')
        base_query += " AND strftime('%Y-%m', v.created_at) = :target_month"
        count_query += " AND strftime('%Y-%m', v.created_at) = :target_month"
        params['target_month'] = target_month
    elif filter == 'year':
        target_year = date or datetime.now().strftime('%Y')
        base_query += " AND strftime('%Y', v.created_at) = :target_year"
        count_query += " AND strftime('%Y', v.created_at) = :target_year"
        params['target_year'] = target_year
    elif filter == 'mastered':
        base_query += " AND s.is_mastered = 1"
        count_query = count_query.replace("FROM vocabulary_items v WHERE", "FROM vocabulary_items v LEFT JOIN vocabulary_srs_states s ON v.id = s.vocab_id WHERE") + " AND s.is_mastered = 1"
    elif filter == 'unmastered':
        base_query += " AND (s.is_mastered = 0 OR s.is_mastered IS NULL)"
        count_query = count_query.replace("FROM vocabulary_items v WHERE", "FROM vocabulary_items v LEFT JOIN vocabulary_srs_states s ON v.id = s.vocab_id WHERE") + " AND (s.is_mastered = 0 OR s.is_mastered IS NULL)"
    elif filter == 'easily_forgotten':
        base_query += " AND v.is_easily_forgotten = 1"
        count_query += " AND v.is_easily_forgotten = 1"

    base_query += " ORDER BY v.created_at DESC LIMIT :limit OFFSET :offset"
    params['limit'] = limit
    params['offset'] = offset

    total_row = db.execute(text(count_query), params).fetchone()
    total_count = int(total_row[0]) if total_row else 0
    
    rows = db.execute(text(base_query), params).fetchall()
    
    items = []
    for r in rows:
        next_rev = None
        if r['next_review_date']:
            try:
                next_rev = datetime.fromisoformat(r['next_review_date'])
            except:
                pass
                
        items.append({
            "id": r['id'],
            "word": r['word'],
            "definition_zh": r['definition_zh'],
            "part_of_speech": r['part_of_speech'],
            "domain": r['domain'],
            "source": r['source'],
            "example_sentence": r['example_sentence'],
            "is_active": bool(r['is_active']),
            "created_at": datetime.fromisoformat(r['created_at']) if r['created_at'] else datetime.now(),
            "traversal_count": r['traversal_count'] or 0,
            "ease_factor": r['ease_factor'] or 2.5,
            "interval_days": r['interval_days'] or 1,
            "next_review_date": next_rev,
            "is_mastered": bool(r['is_mastered']),
            "duplicate_count": r['duplicate_count'] if 'duplicate_count' in r.keys() else 0,
            "is_easily_forgotten": bool(r['is_easily_forgotten']) if 'is_easily_forgotten' in r.keys() else False,
        })

    import math
    return {
        "items": items,
        "total_count": total_count,
        "page": page,
        "total_pages": math.ceil(total_count / limit) if total_count > 0 else 0
    }

@router.post(
    "",
    response_model=VocabItemResponse,
    summary="新增生词",
    description="创建一个新的生词条目并初始化记忆曲线状态。",
)
def create_vocab(item: VocabItemCreate, db = Depends(get_db)):
    # Check if exists — if duplicate, mark as 易忘
    existing = db.execute(text("SELECT id, duplicate_count FROM vocabulary_items WHERE word = :word"), {"word": item.word.lower()}).fetchone()
    if existing:
        new_count = (existing['duplicate_count'] or 0) + 1
        db.execute(text("""
            UPDATE vocabulary_items
            SET duplicate_count = :count, is_easily_forgotten = 1
            WHERE id = :id
        """), {"count": new_count, "id": existing['id']})
        # Increase mastery threshold and reset mastered status
        db.execute(text("""
            UPDATE vocabulary_srs_states
            SET mastery_threshold = mastery_threshold + 2,
                is_mastered = 0,
                next_review_date = :now
            WHERE vocab_id = :vid
        """), {"vid": existing['id'], "now": datetime.now().isoformat()})
        db.commit()
        return {
            "id": existing['id'],
            "word": item.word.lower(),
            "definition_zh": item.definition_zh,
            "part_of_speech": item.part_of_speech,
            "domain": item.domain,
            "source": item.source,
            "example_sentence": item.example_sentence,
            "is_active": True,
            "created_at": datetime.now(),
            "traversal_count": 0,
            "ease_factor": 2.5,
            "interval_days": 1,
            "next_review_date": datetime.now(),
            "is_mastered": False,
            "duplicate": True,
            "duplicate_count": new_count,
        }

    new_id = str(uuid.uuid4())
    now_str = datetime.now().isoformat()
    
    # 1. Insert vocab
    db.execute(text("""
        INSERT INTO vocabulary_items (id, word, definition_zh, part_of_speech, domain, source, example_sentence)
        VALUES (:id, :word, :definition_zh, :part_of_speech, :domain, :source, :example_sentence)
    """), {
        "id": new_id,
        "word": item.word.lower(),
        "definition_zh": item.definition_zh,
        "part_of_speech": item.part_of_speech,
        "domain": item.domain,
        "source": item.source,
        "example_sentence": item.example_sentence
    })
    
    # 2. Insert SRS state
    srs_id = str(uuid.uuid4())
    db.execute(text("""
        INSERT INTO vocabulary_srs_states (id, vocab_id, traversal_count, ease_factor, interval_days, next_review_date, is_mastered, mastery_threshold)
        VALUES (:id, :vocab_id, 0, 2.5, 1, :next_review, 0, 3)
    """), {
        "id": srs_id,
        "vocab_id": new_id,
        "next_review": now_str
    })
    
    db.commit()
    
    return {
        "id": new_id,
        "word": item.word.lower(),
        "definition_zh": item.definition_zh,
        "part_of_speech": item.part_of_speech,
        "domain": item.domain,
        "source": item.source,
        "example_sentence": item.example_sentence,
        "is_active": True,
        "created_at": datetime.now(),
        "traversal_count": 0,
        "ease_factor": 2.5,
        "interval_days": 1,
        "next_review_date": datetime.now(),
        "is_mastered": False
    }


# ══════════════════════════════════════════════
# SCOPE WORDS — 备考范围词表 (MUST be before /{vocab_id} wildcard)
# ══════════════════════════════════════════════

@router.get(
    "/scope-words",
    summary="备考范围词表",
    description="返回用户在备考范围内的词汇掌握情况。",
)
def get_scope_words(
    request: FastAPIRequest,
    page: int = Query(1, ge=1),
    limit: int = Query(30, ge=1, le=100),
    search: str = None,
    status_filter: str = None,
    db=Depends(get_db),
):
    # Get exam type from profile
    exam_type = "cet6"
    try:
        user = getattr(request.state, "user", None)
        if user:
            profile = get_user_profile(user["id"])
            exam_type = profile.get("exam_type", "cet6") or "cet6"
    except Exception:
        pass

    range_row = db.execute(
        "SELECT id, display_name, total_count FROM target_ranges WHERE id = ?",
        (exam_type,)
    ).fetchone()
    range_name = range_row["display_name"] if range_row else exam_type
    target_total = int(range_row["total_count"]) if range_row else 0

    # Build query for all user vocab items with mastery status
    where_parts = ["1=1"]
    params = {}

    if search:
        where_parts.append("v.word LIKE :search")
        params["search"] = f"%{search}%"

    if status_filter == "mastered":
        where_parts.append("s.is_mastered = 1")
    elif status_filter == "unlearned":
        where_parts.append("(s.is_mastered = 0 OR s.is_mastered IS NULL)")

    where_clause = " AND ".join(where_parts)

    # Count
    count_sql = f"""
        SELECT COUNT(*) FROM vocabulary_items v
        LEFT JOIN vocabulary_srs_states s ON s.vocab_id = v.id
        WHERE {where_clause}
    """
    count_row = db.execute(text(count_sql), params).fetchone()
    total = int(count_row[0]) if count_row else 0
    total_pages = max(1, (total + limit - 1) // limit)
    offset = (page - 1) * limit
    params["lim"] = limit
    params["off"] = offset

    rows = db.execute(text(f"""
        SELECT v.id, v.word, v.definition_zh,
               s.is_mastered, s.traversal_count
        FROM vocabulary_items v
        LEFT JOIN vocabulary_srs_states s ON s.vocab_id = v.id
        WHERE {where_clause}
        ORDER BY v.created_at DESC
        LIMIT :lim OFFSET :off
    """), params).fetchall()

    items = []
    for r in rows:
        status = "learned"
        if r["is_mastered"]:
            status = "mastered"
        items.append({
            "word": r["word"],
            "frequency_rank": None,
            "definition_zh": r["definition_zh"] or "",
            "status": status,
            "vocab_id": r["id"],
        })

    # Summary counts
    summary = db.execute(text("""
        SELECT
            COUNT(*) as total,
            SUM(CASE WHEN s.is_mastered = 1 THEN 1 ELSE 0 END) as mastered
        FROM vocabulary_items v
        LEFT JOIN vocabulary_srs_states s ON s.vocab_id = v.id
    """)).fetchone()

    learned_total = int(summary["total"] or 0) if summary else 0
    mastered_total = int(summary["mastered"] or 0) if summary else 0

    return {
        "range_id": exam_type,
        "range_name": range_name,
        "target_total": target_total,
        "total": total,
        "learned": learned_total,
        "mastered": mastered_total,
        "items": items,
        "page": page,
        "total_pages": total_pages,
    }


# ══════════════════════════════════════════════
# AI DOMAIN CLASSIFICATION (MUST be before /{vocab_id} wildcard)
# ══════════════════════════════════════════════

_classify_tasks: Dict[str, Dict[str, Any]] = {}

CLASSIFY_PROMPT = """你是一个专业英语词汇分类专家。请将以下英语词汇按照专业使用场景分类。

每个词可以属于多个专业领域。对每个词，返回它适用的专业领域列表，以及在每个领域下的专业含义。

可选领域包括：
- international-law (国际法)
- commercial-law (商法)
- constitutional-law (宪法学)
- criminal-law (刑法学)
- ip-law (知识产权法)
- financial-law (金融法)
- environmental-law (环境法)
- administrative-law (行政法)
- civil-law (民法)
- procedural-law (诉讼法)
- finance (金融)
- cs (计算机)
- medicine (医学)
- general (通用)

请严格按 JSON 格式返回：
{
  "results": [
    {
      "word": "jurisdiction",
      "domains": [
        {"id": "international-law", "definition": "管辖权，司法管辖范围"},
        {"id": "procedural-law", "definition": "管辖权，法院审理案件的权限"}
      ]
    }
  ]
}

仅返回 JSON，不要额外解释。"""


@router.get(
    "/classify-status",
    summary="分类进度",
)
def get_classify_status(db=Depends(get_db)):
    try:
        total = db.execute(text("SELECT COUNT(*) FROM vocabulary_items")).fetchone()[0] or 0
        classified = db.execute(
            text("SELECT COUNT(DISTINCT vocab_id) FROM vocab_domain_tags")
        ).fetchone()[0] or 0
    except Exception:
        return {"total": 0, "classified": 0, "pending": 0}
    return {
        "total": int(total),
        "classified": int(classified),
        "pending": int(total) - int(classified),
    }


@router.post(
    "/classify",
    summary="AI词汇分类",
)
async def classify_vocab(body: dict = {}, db=Depends(get_db)):
    batch_size = min(body.get("batch_size", 30), 50)

    rows = db.execute(text("""
        SELECT v.id, v.word, v.definition_zh
        FROM vocabulary_items v
        WHERE v.id NOT IN (SELECT DISTINCT vocab_id FROM vocab_domain_tags)
        LIMIT :lim
    """), {"lim": batch_size}).fetchall()

    if not rows:
        return {"status": "done", "classified": 0, "message": "所有词汇已分类完成"}

    words_list = [{"id": r["id"], "word": r["word"], "definition_zh": r["definition_zh"] or ""} for r in rows]
    words_text = "\n".join(f"- {w['word']} ({w['definition_zh']})" for w in words_list)

    try:
        result = await call_claude_json(
            CLASSIFY_PROMPT,
            f"请对以下 {len(words_list)} 个词汇进行专业领域分类：\n{words_text}",
            max_tokens=4000,
            feature_id="vocab_classify",
        )

        classified_count = 0
        ai_results = result.get("results", [])
        word_id_map = {w["word"].lower(): w["id"] for w in words_list}

        for item in ai_results:
            word = item.get("word", "").strip()
            vid = word_id_map.get(word.lower())
            if not vid:
                continue
            for d in item.get("domains", []):
                domain_id = d.get("id", "").strip()
                definition = d.get("definition", "").strip()
                if not domain_id:
                    continue
                try:
                    db.execute(text("""
                        INSERT OR REPLACE INTO vocab_domain_tags (vocab_id, domain, definition_in_domain, classified_by)
                        VALUES (:vid, :dom, :defn, 'ai')
                    """), {"vid": vid, "dom": domain_id, "defn": definition})
                    classified_count += 1
                except Exception:
                    pass

        db.commit()

        total = db.execute(text("SELECT COUNT(*) FROM vocabulary_items")).fetchone()[0] or 0
        done = db.execute(text("SELECT COUNT(DISTINCT vocab_id) FROM vocab_domain_tags")).fetchone()[0] or 0

        return {
            "status": "ok",
            "classified_words": len(ai_results),
            "classified_tags": classified_count,
            "progress": {
                "total": int(total),
                "classified": int(done),
                "pending": int(total) - int(done),
            },
        }
    except Exception as e:
        import logging
        logging.getLogger("xiaotiao").error("Classify error: %s", e)
        raise HTTPException(status_code=500, detail=f"AI 分类失败：{str(e)}")


@router.put(
    "/{vocab_id}",
    response_model=dict,
    summary="更新生词",
    description="更新生词的释义、词性、领域、例句或激活状态。",
)
def update_vocab(vocab_id: str, item: VocabItemUpdate, db = Depends(get_db)):
    updates = []
    params = {"id": vocab_id}
    
    if item.definition_zh is not None:
        updates.append("definition_zh = :def")
        params["def"] = item.definition_zh
    if item.part_of_speech is not None:
        updates.append("part_of_speech = :pos")
        params["pos"] = item.part_of_speech
    if item.domain is not None:
        updates.append("domain = :domain")
        params["domain"] = item.domain
    if item.example_sentence is not None:
        updates.append("example_sentence = :ex")
        params["ex"] = item.example_sentence
    if item.is_active is not None:
        updates.append("is_active = :active")
        params["active"] = 1 if item.is_active else 0
        
    if not updates:
        return {"status": "ok"}
        
    query = f"UPDATE vocabulary_items SET {', '.join(updates)} WHERE id = :id"
    res = db.execute(text(query), params)
    if res.rowcount == 0:
        raise HTTPException(status_code=404, detail="未找到该生词。")
    db.commit()
    
    return {"status": "ok"}

@router.delete(
    "/{vocab_id}",
    summary="删除生词",
    description="从生词本中移除指定生词。",
)
def delete_vocab(vocab_id: str, db = Depends(get_db)):
    res = db.execute(text("DELETE FROM vocabulary_items WHERE id = :id"), {"id": vocab_id})
    if res.rowcount == 0:
        raise HTTPException(status_code=404, detail="未找到该生词。")
    db.commit()
    return {"status": "ok"}


@router.put(
    "/batch-mastery",
    summary="批量切换掌握状态",
    description="批量标记或取消标记多个生词为已掌握。",
)
def batch_toggle_mastery(body: dict, db = Depends(get_db)):
    ids = body.get("ids", [])
    mastered = body.get("is_mastered", True)
    if not ids:
        raise HTTPException(status_code=400, detail="请提供词汇ID列表。")
    updated = 0
    for vid in ids:
        srs = db.execute(text("SELECT id FROM vocabulary_srs_states WHERE vocab_id = :vid"), {"vid": vid}).fetchone()
        if srs:
            db.execute(text("UPDATE vocabulary_srs_states SET is_mastered = :m WHERE vocab_id = :vid"),
                       {"m": 1 if mastered else 0, "vid": vid})
        else:
            srs_id = str(uuid.uuid4())
            db.execute(text("""
                INSERT INTO vocabulary_srs_states (id, vocab_id, traversal_count, ease_factor, interval_days, next_review_date, is_mastered)
                VALUES (:id, :vid, 0, 2.5, 1, :now, :m)
            """), {"id": srs_id, "vid": vid, "now": datetime.now().isoformat(), "m": 1 if mastered else 0})
        updated += 1
    db.commit()
    return {"status": "ok", "updated": updated, "is_mastered": mastered}


@router.put(
    "/{vocab_id}/mastery",
    summary="切换掌握状态",
    description="标记或取消标记某生词为已掌握。",
)
def toggle_mastery(vocab_id: str, body: dict, db = Depends(get_db)):
    mastered = body.get("is_mastered", False)
    # Check if SRS state exists
    srs = db.execute(text("SELECT id FROM vocabulary_srs_states WHERE vocab_id = :vid"), {"vid": vocab_id}).fetchone()
    if srs:
        db.execute(text("""
            UPDATE vocabulary_srs_states SET is_mastered = :m WHERE vocab_id = :vid
        """), {"m": 1 if mastered else 0, "vid": vocab_id})
    else:
        srs_id = str(uuid.uuid4())
        db.execute(text("""
            INSERT INTO vocabulary_srs_states (id, vocab_id, traversal_count, ease_factor, interval_days, next_review_date, is_mastered)
            VALUES (:id, :vid, 0, 2.5, 1, :now, :m)
        """), {"id": srs_id, "vid": vocab_id, "now": datetime.now().isoformat(), "m": 1 if mastered else 0})
    db.commit()
    return {"status": "ok", "is_mastered": mastered}


CONCEPT_ANALYSIS_PROMPT = """You are a bilingual English-Chinese vocabulary expert. Given an English word or short phrase, provide a clear, educational analysis.

Return a JSON object with this structure:
{
  "word": "the word",
  "phonetic": "phonetic transcription",
  "part_of_speech": "n./v./adj./adv./phrase",
  "definition_en": "English definition",
  "definition_zh": "中文释义",
  "explanation": "Detailed Chinese explanation of the concept, including usage context and nuances (2-3 sentences)",
  "examples": [
    {"en": "Example sentence in English", "zh": "例句中文翻译"}
  ],
  "related_terms": [
    {"term": "related word", "zh": "释义"}
  ]
}

Rules:
- Keep definitions concise but informative
- Provide 2-3 example sentences
- Provide 2-4 related terms
- Explanation should be in Chinese, detailed and educational
- Return ONLY the JSON object, no markdown wrappers
"""


@router.post(
    "/concept",
    summary="概念解析",
    description="对单个英文词汇或短语进行语义解析，返回释义、例句和关联词汇。",
)
async def concept_analysis(req: dict):
    """Analyze a single English word/phrase and return educational content."""
    word = req.get("word", "").strip()
    if not word:
        raise HTTPException(status_code=400, detail="请提供要解析的词汇。")
    if len(word) > 100:
        raise HTTPException(status_code=400, detail="词汇过长。")

    try:
        data = await call_claude_json(
            CONCEPT_ANALYSIS_PROMPT,
            f"Please analyze this English word/phrase: {word}",
            max_tokens=2000,
            feature_id="concept_analysis",
        )
        return data
    except Exception as e:
        import logging
        logging.getLogger("xiaotiao").error("Concept analysis error: %s", e)
        raise HTTPException(status_code=500, detail=f"AI 解析失败：{str(e)}")


# ── 异步任务系统 ──────────────────────────────────
# 内存中存储任务状态（单机足够，无需 Redis）
_import_tasks: Dict[str, Dict[str, Any]] = {}
_CHUNK_SIZE = 5000  # 每个 chunk 最大字符数
_MAX_CHUNKS = 20    # 最大 chunk 数量


def _split_text_into_chunks(text: str, chunk_size: int = _CHUNK_SIZE) -> List[str]:
    """将长文本按段落边界分割成多个 chunk。"""
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    lines = text.split('\n')
    current_chunk = []
    current_len = 0

    for line in lines:
        line_len = len(line) + 1  # +1 for newline
        if current_len + line_len > chunk_size and current_chunk:
            chunks.append('\n'.join(current_chunk))
            current_chunk = [line]
            current_len = line_len
        else:
            current_chunk.append(line)
            current_len += line_len

    if current_chunk:
        chunks.append('\n'.join(current_chunk))

    # 限制最大 chunk 数
    return chunks[:_MAX_CHUNKS]


def _extract_file_content(contents: bytes, filename: str) -> str:
    """从文件中提取文本内容（同步，不涉及 LLM）。"""
    extracted_text = ""

    if filename.endswith((".txt", ".md")):
        try:
            extracted_text = contents.decode("utf-8")
        except UnicodeDecodeError:
            extracted_text = contents.decode("gbk", errors="ignore")

    elif filename.endswith(".csv"):
        try:
            text_content = contents.decode("utf-8")
        except UnicodeDecodeError:
            text_content = contents.decode("gbk", errors="ignore")
        try:
            reader = csv.reader(StringIO(text_content))
            rows = []
            for row in reader:
                non_empty = [cell.strip() for cell in row if cell.strip()]
                if non_empty:
                    rows.append(" | ".join(non_empty))
            extracted_text = "\n".join(rows)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"解析 CSV 失败：{str(e)}")

    elif filename.endswith((".xlsx", ".xls")):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(BytesIO(contents), data_only=True)
            rows = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                rows.append(f"--- Sheet: {sheet_name} ---")
                for row in sheet.iter_rows(values_only=True):
                    non_empty = [str(cell).strip() for cell in row if cell is not None]
                    if non_empty:
                        rows.append(" | ".join(non_empty))
            extracted_text = "\n".join(rows)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"解析 Excel 失败：{str(e)}")

    elif filename.endswith((".docx", ".doc")):
        try:
            import docx
            doc = docx.Document(BytesIO(contents))
            extracted_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"解析 Word 文档失败：{str(e)}")

    return extracted_text


async def _process_chunk(chunk: str, domain: str, chunk_index: int, total_chunks: int) -> List[dict]:
    """用 LLM 处理单个文本块，提取词汇。"""
    user_prompt = (
        f"Domain focus: {domain}\n"
        f"This is chunk {chunk_index + 1} of {total_chunks}.\n\n"
        f"Please extract English vocabulary words from the following content:\n\n{chunk}"
    )
    try:
        data = await call_claude_json(
            VOCAB_IMPORT_SYSTEM_PROMPT, user_prompt,
            max_tokens=4000, feature_id="vocab_import"
        )
        words = data.get("words", [])
        return [
            {
                "word": w.get("word", ""),
                "definition_zh": w.get("definition_zh", ""),
                "part_of_speech": w.get("part_of_speech", "n."),
                "example_sentence": w.get("example_sentence", ""),
            }
            for w in words if w.get("word", "").strip()
        ]
    except Exception as e:
        import logging
        logging.getLogger("xiaotiao").warning("Chunk %d extraction failed: %s", chunk_index, e)
        return []


def _deduplicate_words(all_words: List[dict]) -> List[dict]:
    """按 word 去重，保留第一次出现的条目。"""
    seen = set()
    result = []
    for w in all_words:
        key = w["word"].lower().strip()
        if key and key not in seen:
            seen.add(key)
            result.append(w)
    return result


async def _run_import_task(task_id: str, extracted_text: str, domain: str):
    """后台任务：分块处理文本并提取词汇。"""
    task = _import_tasks[task_id]
    try:
        chunks = _split_text_into_chunks(extracted_text)
        task["total_chunks"] = len(chunks)
        task["status"] = "processing"
        task["message"] = f"正在提取词汇 (0/{len(chunks)})"

        all_words = []
        for i, chunk in enumerate(chunks):
            task["current_chunk"] = i + 1
            task["progress"] = int((i / len(chunks)) * 100)
            task["message"] = f"正在提取词汇 ({i + 1}/{len(chunks)})"

            words = await _process_chunk(chunk, domain, i, len(chunks))
            all_words.extend(words)

        # 去重
        unique_words = _deduplicate_words(all_words)
        task["status"] = "done"
        task["progress"] = 100
        task["words"] = unique_words
        task["message"] = f"提取完成，共 {len(unique_words)} 个词汇"
    except Exception as e:
        task["status"] = "error"
        task["error"] = str(e)
        task["message"] = f"提取失败：{str(e)}"


async def _run_image_import_task(task_id: str, base64_image: str, media_type: str, domain: str):
    """后台任务：图片识别词汇。"""
    task = _import_tasks[task_id]
    try:
        task["status"] = "processing"
        task["progress"] = 30
        task["message"] = "AI 正在识别图片中的词汇..."

        data = await call_claude_vision_json(
            system_prompt=VOCAB_IMPORT_SYSTEM_PROMPT,
            user_prompt=f"Domain focus: {domain}\n\nPlease extract all English vocabulary words visible in this image.",
            base64_image=base64_image,
            media_type=media_type,
            max_tokens=4000,
            feature_id="vocab_import",
        )
        words = data.get("words", [])
        result = [
            {
                "word": w.get("word", ""),
                "definition_zh": w.get("definition_zh", ""),
                "part_of_speech": w.get("part_of_speech", "n."),
                "example_sentence": w.get("example_sentence", ""),
            }
            for w in words if w.get("word", "").strip()
        ]
        task["status"] = "done"
        task["progress"] = 100
        task["words"] = result
        task["message"] = f"提取完成，共 {len(result)} 个词汇"
    except Exception as e:
        task["status"] = "error"
        task["error"] = str(e)
        task["message"] = f"图片识别失败：{str(e)}"


@router.post(
    "/import-file",
    summary="导入词汇文件（异步）",
    description="上传文件后立即返回 task_id，前端轮询 /vocab/import-task/{task_id} 获取进度。",
)
async def import_vocab_file(
    file: UploadFile = File(...),
    domain: str = Form("general"),
):
    """Parse uploaded file and start async extraction task."""
    contents = await file.read()
    filename = (file.filename or "unknown").lower()

    task_id = str(uuid.uuid4())
    _import_tasks[task_id] = {
        "status": "pending",
        "progress": 0,
        "current_chunk": 0,
        "total_chunks": 1,
        "words": [],
        "error": None,
        "message": "正在准备...",
        "created_at": datetime.now().isoformat(),
    }

    # 图片文件：走 vision 模型
    if filename.endswith((".png", ".jpg", ".jpeg")):
        base64_image = base64.b64encode(contents).decode("utf-8")
        media_type = "image/png" if filename.endswith(".png") else "image/jpeg"
        asyncio.create_task(_run_image_import_task(task_id, base64_image, media_type, domain))
        return {"task_id": task_id}

    # 文本类文件：提取文本内容
    if not filename.endswith((".txt", ".md", ".csv", ".xlsx", ".xls", ".docx", ".doc")):
        raise HTTPException(
            status_code=400,
            detail="不支持的文件格式。支持：.txt, .md, .csv, .xlsx, .xls, .docx, .doc, .png, .jpg, .jpeg",
        )

    extracted_text = _extract_file_content(contents, filename)
    if not extracted_text.strip():
        raise HTTPException(status_code=400, detail="文件内容为空。")

    _import_tasks[task_id]["message"] = f"文件已解析，共 {len(extracted_text)} 字符，开始提取..."
    asyncio.create_task(_run_import_task(task_id, extracted_text, domain))
    return {"task_id": task_id}


@router.get(
    "/import-task/{task_id}",
    summary="查询导入任务状态",
    description="前端轮询此端点获取异步导入任务的进度和结果。",
)
async def get_import_task_status(task_id: str):
    """Poll import task status."""
    task = _import_tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在或已过期。")

    response = {
        "status": task["status"],
        "progress": task["progress"],
        "current_chunk": task["current_chunk"],
        "total_chunks": task["total_chunks"],
        "message": task["message"],
    }

    if task["status"] == "done":
        response["words"] = task["words"]
        # 清理已完成的任务（延迟 60 秒）
        asyncio.get_event_loop().call_later(60, lambda: _import_tasks.pop(task_id, None))

    if task["status"] == "error":
        response["error"] = task["error"]
        asyncio.get_event_loop().call_later(60, lambda: _import_tasks.pop(task_id, None))

    return response


@router.post(
    "/batch",
    summary="批量导入词汇",
    description="批量添加多个词汇到生词本。",
)
def batch_create_vocab(
    items: List[VocabItemCreate],
    db=Depends(get_db),
):
    """Import multiple words at once, skipping duplicates."""
    imported = 0
    skipped = 0
    for item in items:
        existing = db.execute(
            text("SELECT id FROM vocabulary_items WHERE word = :word"),
            {"word": item.word.lower()},
        ).fetchone()
        if existing:
            skipped += 1
            continue

        new_id = str(uuid.uuid4())
        db.execute(
            text("""
                INSERT INTO vocabulary_items (id, word, definition_zh, part_of_speech, domain, source, example_sentence)
                VALUES (:id, :word, :definition_zh, :part_of_speech, :domain, :source, :example_sentence)
            """),
            {
                "id": new_id,
                "word": item.word.lower(),
                "definition_zh": item.definition_zh,
                "part_of_speech": item.part_of_speech,
                "domain": item.domain,
                "source": item.source or "file_import",
                "example_sentence": item.example_sentence,
            },
        )
        srs_id = str(uuid.uuid4())
        db.execute(
            text("""
                INSERT INTO vocabulary_srs_states (id, vocab_id, traversal_count, ease_factor, interval_days, next_review_date, is_mastered)
                VALUES (:id, :vocab_id, 0, 2.5, 1, :next_review, 0)
            """),
            {
                "id": srs_id,
                "vocab_id": new_id,
                "next_review": datetime.now().isoformat(),
            },
        )
        imported += 1

    db.commit()
    return {"imported": imported, "skipped": skipped, "total": len(items)}


@router.get(
    "/scope-stats",
    summary="备考范围词汇统计",
    description="返回用户当前备考范围内的词汇总数、已学习数和已掌握数。",
)
def get_scope_stats(request: FastAPIRequest, db=Depends(get_db)):
    # Get user exam_type from profile
    exam_type = "cet6"  # default
    try:
        user = getattr(request.state, "user", None)
        if user:
            profile = get_user_profile(user["id"])
            exam_type = profile.get("exam_type", "cet6") or "cet6"
    except Exception:
        pass

    # Get target range info
    range_row = db.execute(
        "SELECT id, display_name, total_count, description FROM target_ranges WHERE id = ?",
        (exam_type,)
    ).fetchone()

    if not range_row:
        return {
            "scope_id": exam_type,
            "scope_name": exam_type,
            "total": 0,
            "learned": 0,
            "mastered": 0,
            "description": "",
        }

    total = range_row["total_count"] or 0
    scope_name = range_row["display_name"]
    description = range_row["description"] or ""

    # Count how many scope words are in user's vocabulary
    learned = 0
    mastered = 0
    try:
        learned_row = db.execute(
            "SELECT COUNT(*) FROM vocabulary_items"
        ).fetchone()
        learned = int(learned_row[0]) if learned_row and learned_row[0] else 0

        mastered_row = db.execute("""
            SELECT COUNT(*) FROM vocabulary_items v
            INNER JOIN vocabulary_srs_states s ON s.vocab_id = v.id AND s.is_mastered = 1
        """).fetchone()
        mastered = int(mastered_row[0]) if mastered_row and mastered_row[0] else 0
    except Exception:
        pass

    return {
        "scope_id": exam_type,
        "scope_name": scope_name,
        "total": total,
        "learned": learned,
        "mastered": mastered,
        "description": description,
    }


@router.get(
    "/stats",
    response_model=VocabStatsResponse,
    summary="生词统计",
    description="获取生词总量、活跃数与今日需复习数量。",
)
def get_stats(db = Depends(get_db)):
    total_row = db.execute(text("SELECT count(1) FROM vocabulary_items")).fetchone()
    active_row = db.execute(text("SELECT count(1) FROM vocabulary_items WHERE is_active = 1")).fetchone()
    total = int(total_row[0]) if total_row else 0
    active = int(active_row[0]) if active_row else 0
    
    srs_stats = db.execute(text("""
        SELECT 
            SUM(CASE WHEN is_mastered = 1 THEN 1 ELSE 0 END) as mastered,
            SUM(CASE WHEN is_mastered = 0 AND next_review_date <= :now THEN 1 ELSE 0 END) as due_today
        FROM vocabulary_srs_states
    """), {"now": datetime.now().isoformat()}).fetchone()

    # Time-based counts
    today_str = datetime.now().strftime('%Y-%m-%d')
    month_str = datetime.now().strftime('%Y-%m')
    year_str = datetime.now().strftime('%Y')

    today_count = db.execute(text(
        "SELECT count(1) FROM vocabulary_items WHERE date(created_at) = :d"
    ), {"d": today_str}).fetchone()[0] or 0
    month_count = db.execute(text(
        "SELECT count(1) FROM vocabulary_items WHERE strftime('%Y-%m', created_at) = :m"
    ), {"m": month_str}).fetchone()[0] or 0
    year_count = db.execute(text(
        "SELECT count(1) FROM vocabulary_items WHERE strftime('%Y', created_at) = :y"
    ), {"y": year_str}).fetchone()[0] or 0
    easily_forgotten = db.execute(text(
        "SELECT count(1) FROM vocabulary_items WHERE is_easily_forgotten = 1"
    )).fetchone()[0] or 0
    
    return {
        "total": total or 0,
        "active": active or 0,
        "mastered": srs_stats[0] if srs_stats and srs_stats[0] else 0,
        "need_review_today": srs_stats[1] if srs_stats and srs_stats[1] else 0,
        "today_count": int(today_count),
        "month_count": int(month_count),
        "year_count": int(year_count),
        "easily_forgotten": int(easily_forgotten),
    }


@router.get(
    "/export",
    summary="导出词汇表",
    description="导出当前筛选条件下的词汇为 Word 文档。",
)
def export_vocab(
    db=Depends(get_db),
    filter: str = None,
    date: str = None,
    search: str = None,
    domain: str = None,
):
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT

    # Build query (reuse filter logic)
    base_query = """
        SELECT v.*, s.is_mastered, v.is_easily_forgotten
        FROM vocabulary_items v
        LEFT JOIN vocabulary_srs_states s ON v.id = s.vocab_id
        WHERE 1=1
    """
    params = {}
    if search:
        base_query += " AND v.word LIKE :search"
        params['search'] = f"%{search}%"
    if domain:
        base_query += " AND v.domain = :domain"
        params['domain'] = domain

    title_suffix = "全部生词"
    if filter == 'today':
        target_date = date or datetime.now().strftime('%Y-%m-%d')
        base_query += " AND date(v.created_at) = :target_date"
        params['target_date'] = target_date
        title_suffix = f"当日生词本 ({target_date})"
    elif filter == 'month':
        target_month = date or datetime.now().strftime('%Y-%m')
        base_query += " AND strftime('%Y-%m', v.created_at) = :target_month"
        params['target_month'] = target_month
        title_suffix = f"当月生词本 ({target_month})"
    elif filter == 'year':
        target_year = date or datetime.now().strftime('%Y')
        base_query += " AND strftime('%Y', v.created_at) = :target_year"
        params['target_year'] = target_year
        title_suffix = f"当年生词本 ({target_year})"
    elif filter == 'mastered':
        base_query += " AND s.is_mastered = 1"
        title_suffix = "已掌握词汇"
    elif filter == 'unmastered':
        base_query += " AND (s.is_mastered = 0 OR s.is_mastered IS NULL)"
        title_suffix = "未掌握词汇"
    elif filter == 'easily_forgotten':
        base_query += " AND v.is_easily_forgotten = 1"
        title_suffix = "易忘生词"

    base_query += " ORDER BY v.created_at DESC"
    rows = db.execute(text(base_query), params).fetchall()

    # Create Word document
    doc = docx.Document()
    doc.add_heading(f"生词本 — {title_suffix}", level=1)
    doc.add_paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}　共 {len(rows)} 个词汇")

    if rows:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['单词', '中文释义', '词性', '专业领域', '状态']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        for r in rows:
            status = '已掌握' if r['is_mastered'] else ('易忘' if r['is_easily_forgotten'] else '学习中')
            row_cells = table.add_row().cells
            row_cells[0].text = r['word'] or ''
            row_cells[1].text = r['definition_zh'] or ''
            row_cells[2].text = r['part_of_speech'] or ''
            row_cells[3].text = r['domain'] or ''
            row_cells[4].text = status

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    # Use RFC 5987 encoding for non-ASCII filenames
    from urllib.parse import quote
    safe_name = title_suffix.replace(' ', '_')
    encoded_name = quote(f"vocab_{safe_name}.docx")

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=vocab_export.docx; filename*=UTF-8''{encoded_name}"}
    )

